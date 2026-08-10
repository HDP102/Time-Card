"""
ZIA Comprehensive Policy Analysis
-------------------------------------
Combines URL filtering and SSL inspection cleanup strategy with
deep category content analysis for thorough recommendations.

Reads from same folder:
  - url_filtering.json    (required)
  - sslpol.json           (required)
  - url_cat.json          (required — custom URL category data)
  - *_WEB_log.csv         (optional — Web Insights exports for hit counts)

PowerShell:
    pip install openpyxl python-docx
    python zia_comprehensive_analysis.py url_filtering.json sslpol.json

Output (in current directory):
    - ZIA_URL_Filtering_Cleanup_Strategy.xlsx
    - ZIA_SSL_Inspection_Cleanup_Strategy.xlsx
    - ZIA_URL_Filtering_Proposal.docx
    - ZIA_SSL_Inspection_Proposal.docx
"""

import json, sys, os, re, glob
from datetime import datetime
from collections import Counter, defaultdict

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
except ImportError:
    print("ERROR: openpyxl required. Install: pip install openpyxl"); sys.exit(1)

try:
    from docx import Document as DocxDoc
    from docx.shared import Pt
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("NOTE: python-docx not installed — Word docs skipped. Install: pip install python-docx")


# === STYLES ===
URL_STYLES = {
    "hdr_fill": PatternFill("solid", fgColor="4472C4"),
    "hdr_font": Font(name="Calibri", bold=True, size=11, color="FFFFFF"),
    "norm": Font(name="Calibri", size=11),
    "bold": Font(name="Calibri", bold=True, size=11),
    "title": Font(name="Calibri", bold=True, size=14),
    "subtitle": Font(name="Calibri", bold=True, size=12),
    "red": Font(name="Calibri", size=11, color="CC0000"),
    "bdr": Border(left=Side(style="thin", color="B4B4B4"), right=Side(style="thin", color="B4B4B4"),
                  top=Side(style="thin", color="B4B4B4"), bottom=Side(style="thin", color="B4B4B4")),
}
SSL_STYLES = {
    "hdr_fill": PatternFill("solid", fgColor="404040"),
    "hdr_font": Font(name="Arial", bold=True, size=10, color="FFFFFF"),
    "norm": Font(name="Arial", size=10),
    "bold": Font(name="Arial", bold=True, size=10),
    "title": Font(name="Arial", bold=True, size=13),
    "subtitle": Font(name="Arial", bold=True, size=11),
    "red": Font(name="Arial", size=10, color="CC0000"),
    "bdr": Border(left=Side(style="thin", color="AAAAAA"), right=Side(style="thin", color="AAAAAA"),
                  top=Side(style="thin", color="AAAAAA"), bottom=Side(style="thin", color="AAAAAA")),
}
L_RED = PatternFill("solid", fgColor="F2DCDB")
L_ORANGE = PatternFill("solid", fgColor="FDE9D9")
L_YELLOW = PatternFill("solid", fgColor="FFFFCC")
L_GREEN = PatternFill("solid", fgColor="D8E4BC")

def hdr(ws, row, headers, sty):
    for c, h in enumerate(headers, 1):
        cl = ws.cell(row=row, column=c, value=h)
        cl.font = sty["hdr_font"]; cl.fill = sty["hdr_fill"]; cl.border = sty["bdr"]
        cl.alignment = Alignment(wrap_text=True, vertical="center")

def cel(ws, row, col, val, sty, font=None, fill=None):
    c = ws.cell(row=row, column=col, value=val)
    c.font = font or sty["norm"]; c.border = sty["bdr"]
    if fill: c.fill = fill
    c.alignment = Alignment(wrap_text=True, vertical="top")
    return c

def sw(ws, widths):
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

def trunc(items, n, sep=", "):
    if not items: return ""
    s = sep.join(str(x) for x in items[:n])
    if len(items) > n: s += f"... (+{len(items)-n} more)"
    return s


# === DATA LOADING ===
def load_categories(folder):
    for name in ["url_cat.json", "zia_custom_categories.json"]:
        p = os.path.join(folder, name)
        if os.path.exists(p):
            with open(p, "r", encoding="utf-8") as f: data = json.load(f)
            cats = data.get("categories", data) if isinstance(data, dict) else data
            return {c["id"]: c for c in cats}, cats
    return None, None

def resolve(cid, cat_map):
    if cid.startswith("CUSTOM_") and cid in cat_map:
        return f"{cat_map[cid].get('name', cid)} ({cid})"
    return cid

def load_web_logs(folder):
    csvs = glob.glob(os.path.join(folder, "*WEB_log*")) + glob.glob(os.path.join(folder, "*web_log*"))
    csvs = list(set(csvs))
    url_hits, ssl_hits = Counter(), Counter()
    for csv_file in csvs:
        try:
            with open(csv_file, "rb") as f: raw = f.read(500)
            enc = "utf-16-le" if raw[:2]==b"\xff\xfe" else ("utf-8-sig" if raw[:3]==b"\xef\xbb\xbf" else "utf-8")
            with open(csv_file, "r", encoding=enc, errors="replace") as f: lines = f.readlines()
            hi = 0
            for i, line in enumerate(lines):
                if any(kw in line.lower() for kw in ["event time","policy action","ssl policy"]):
                    hi = i; break
                if i > 15: break
            sep = "\t" if lines[hi].count("\t") > lines[hi].count(",") else ","
            headers = [h.strip().strip('"') for h in lines[hi].split(sep)]
            url_col = ssl_col = None
            for idx, h in enumerate(headers):
                hl = h.lower()
                if "url filtering policy" in hl: url_col = idx
                elif "ssl policy" in hl: ssl_col = idx
            for line in lines[hi+1:]:
                fields = [f.strip().strip('"') for f in line.split(sep)]
                if url_col is not None and url_col < len(fields):
                    v = fields[url_col]
                    if v and v.lower() not in ("none","","na"): url_hits[v] += 1
                if ssl_col is not None and ssl_col < len(fields):
                    v = fields[ssl_col]
                    if v and v.lower() not in ("none","","na"): ssl_hits[v] += 1
            print(f"    {os.path.basename(csv_file)}: processed")
        except Exception as e:
            print(f"    {os.path.basename(csv_file)}: ERROR {e}")
    return dict(url_hits), dict(ssl_hits)


# === ANALYSIS ===
def analyze_all(url_rules, ssl_rules, cat_map, cat_list):
    f = {}
    sorted_url = sorted(url_rules, key=lambda x: x.get("order", 999))
    sorted_ssl = sorted(ssl_rules, key=lambda x: x.get("order", 999))

    # Usage maps
    uu, su, std = defaultdict(list), defaultdict(list), set()
    for r in url_rules:
        for c in r.get("urlCategories", []):
            if c.startswith("CUSTOM_"): uu[c].append(r["name"])
            else: std.add(c)
    for r in ssl_rules:
        for c in r.get("urlCategories", []):
            if c.startswith("CUSTOM_"): su[c].append(r["name"])
    f["url_usage"], f["ssl_usage"], f["std_cats"] = dict(uu), dict(su), sorted(std)

    # Duplicate URLs
    u2c = defaultdict(list)
    for c in cat_list:
        for u in c.get("urls", []):
            u2c[u].append({"id": c["id"], "name": c.get("name","")})
    f["duplicates"] = {u: cats for u, cats in u2c.items() if len(cats) > 1}

    # Cross-policy shared categories
    shared = set(uu.keys()) & set(su.keys())
    f["shared_cats"] = []
    for cid in sorted(shared):
        cat = cat_map.get(cid, {})
        f["shared_cats"].append({
            "id": cid, "name": cat.get("name",""),
            "urls": cat.get("urls",[])[:10], "url_count": len(cat.get("urls",[])),
            "url_rules": uu[cid], "ssl_rules": su[cid],
        })

    # Broad wildcards in SSL bypass
    bypass_cats = set()
    for r in sorted_ssl:
        if r["action"]["type"] == "DO_NOT_DECRYPT":
            for c in r.get("urlCategories", []): bypass_cats.add(c)
    broad_set = {".blob.core.windows.net",".amazonaws.com",".azure.com",".microsoft.com",
                 ".windows.net",".office.com",".google.com",".github.com"}
    f["broad_wildcards"] = []
    for cid in bypass_cats:
        cat = cat_map.get(cid, {})
        for u in cat.get("urls", []):
            if u in broad_set:
                f["broad_wildcards"].append({"cat_id": cid, "cat_name": cat.get("name",""), "url": u})

    # AI URLs
    f["ai_urls"] = []
    for c in cat_list:
        for u in c.get("urls", []):
            if any(kw in u.lower() for kw in ["openai","chatgpt","claude","anthropic","gemini"]):
                f["ai_urls"].append({"url": u, "cat_id": c["id"], "cat_name": c.get("name","")})

    # DLP broad wildcards
    dlp = cat_map.get("CUSTOM_13", {})
    f["dlp_broad"] = [u for u in dlp.get("urls", []) if u.startswith(".") and u.count(".") <= 2]

    # Blocked/allowed conflicts
    blocked = set(cat_map.get("CUSTOM_02",{}).get("urls",[]) + cat_map.get("CUSTOM_05",{}).get("urls",[]))
    allowed = set(cat_map.get("CUSTOM_04",{}).get("urls",[]) + cat_map.get("CUSTOM_07",{}).get("urls",[]) +
                  cat_map.get("CUSTOM_03",{}).get("urls",[]))
    f["block_allow"] = sorted(blocked & allowed)

    # Content mismatches
    f["content_mix"] = []
    for c in cat_list:
        urls = c.get("urls", [])
        name = c.get("name","").lower()
        if len(urls) < 3: continue
        ms = sum(1 for u in urls if "microsoft" in u.lower() or "azure" in u.lower())
        if ms > 3:
            kws = ["gitlab","splunk","informatica","dell","cisco","axonius","sailpoint","netsec","mulesoft","imaging"]
            if any(kw in name for kw in kws):
                f["content_mix"].append({"id": c["id"], "name": c.get("name",""), "ms_count": ms, "total": len(urls)})

    # Rule shadowing
    f["shadows"] = []
    blocks = [(r["order"], r) for r in sorted_url if r.get("action") == "BLOCK"]
    allows = [(r["order"], r) for r in sorted_url if r.get("action") == "ALLOW"]
    for bo, br in blocks:
        bc = set(br.get("urlCategories", []))
        if not bc: continue
        for ao, ar in allows:
            if ao < bo:
                ov = set(ar.get("urlCategories", [])) & bc
                if ov and len(ov) >= 3:
                    f["shadows"].append({"allow": ar["name"], "a_order": ao, "block": br["name"],
                                        "b_order": bo, "count": len(ov), "cats": trunc(sorted(ov), 8)})

    # Merge opportunities
    sg = defaultdict(list)
    for r in sorted_url:
        locs = tuple(sorted(l["name"] for l in r.get("locations", [])))
        if locs:
            key = (locs, tuple(sorted(d["name"] for d in r.get("departments",[]))),
                   tuple(sorted(u["name"] for u in r.get("users",[]))), r.get("action",""))
            sg[key].append(r)
    f["merges"] = [{"rules": [r["name"] for r in rules], "location": ", ".join(key[0])[:80],
                    "action": key[3], "count": len(rules)}
                   for key, rules in sg.items() if len(rules) > 1]

    # Unused categories
    f["unused"] = [{"id": c["id"], "name": c.get("name",""), "urls": len(c.get("urls",[]))}
                   for c in cat_list if c["id"] not in uu and c["id"] not in su]

    # Test items
    f["test_items"] = []
    for r in url_rules:
        if any(kw in r["name"].lower() for kw in ["test","temp","poc"]):
            f["test_items"].append({"name": r["name"], "type": "URL rule", "order": r.get("order")})
    for r in ssl_rules:
        if any(kw in r["name"].lower() for kw in ["test","temp","poc"]):
            f["test_items"].append({"name": r["name"], "type": "SSL rule", "order": r.get("order")})

    # SSL label stats
    ls = defaultdict(lambda: {"decrypt":0,"bypass":0,"bypass_all":0,"bypass_eval":0})
    for r in sorted_ssl:
        for l in r.get("labels", []):
            if r["action"]["type"] == "DECRYPT": ls[l["name"]]["decrypt"] += 1
            else:
                ls[l["name"]]["bypass"] += 1
                bp = r["action"].get("doNotDecryptSubActions",{}).get("bypassOtherPolicies","")
                if bp == True: ls[l["name"]]["bypass_all"] += 1
                elif bp == False: ls[l["name"]]["bypass_eval"] += 1
    f["label_stats"] = dict(ls)

    # DECRYPT coverage
    f["decrypt_rules"] = [{"name": r["name"], "order": r.get("order"),
        "labels": ", ".join(l["name"] for l in r.get("labels",[])) or "(none)",
        "groups": ", ".join(g["name"] for g in r.get("groups",[])) or "(none)",
        "predefined": r.get("predefined", False)}
        for r in sorted_ssl if r["action"]["type"] == "DECRYPT"]

    # E-sign
    f["esign"] = []
    for c in cat_list:
        for u in c.get("urls", []):
            for t in ["signnow","signwell","hellosign","assuresign","insuresign","docusign","pandadoc"]:
                if t in u.lower():
                    f["esign"].append({"url": u, "cat": c.get("name",""), "cat_id": c["id"]})

    # Blocked sites audit
    f["blocked_sites"] = cat_map.get("CUSTOM_02",{}).get("urls",[])

    # Stats
    f["single_url"] = sum(1 for c in cat_list if len(c.get("urls",[])) == 1)
    f["large_cats"] = sum(1 for c in cat_list if len(c.get("urls",[])) > 100)

    return f


# === URL FILTERING EXCEL ===
def build_url_excel(url_rules, findings, cat_map, url_hits):
    wb = Workbook(); sty = URL_STYLES
    sorted_rules = sorted(url_rules, key=lambda x: x.get("order", 999))

    # --- Summary ---
    ws = wb.active; ws.title = "Summary"
    ws["A1"] = "ZIA URL Filtering Rules — Cleanup & Analysis"; ws["A1"].font = sty["title"]
    r = 3
    for label, val in [
        ("Total rules", len(url_rules)), ("ALLOW", sum(1 for x in url_rules if x.get("action")=="ALLOW")),
        ("BLOCK", sum(1 for x in url_rules if x.get("action")=="BLOCK")),
        ("ISOLATE", sum(1 for x in url_rules if x.get("action")=="ISOLATE")),
        ("Disabled", sum(1 for x in url_rules if x.get("state")=="DISABLED")),
        ("",""), ("Security Findings",""),
        ("Rule shadowing candidates", len(findings["shadows"])),
        ("Cross-policy category mismatches", sum(1 for c in findings["shared_cats"] if c["id"] in ("CUSTOM_88","CUSTOM_247","CUSTOM_161","CUSTOM_25"))),
        ("Categories with mixed content", len(findings["content_mix"])),
        ("Blocked/allowed URL conflicts", len(findings["block_allow"])),
        ("",""), ("Hygiene",""),
        ("Merge opportunities (identical scope)", len(findings["merges"])),
        ("Unused categories", len(findings["unused"])),
        ("Test items in production", len(findings["test_items"])),
        ("Duplicate URLs across categories", len(findings["duplicates"])),
    ]:
        if val == "":
            ws.cell(row=r, column=1, value=label).font = sty["subtitle"] if label else sty["norm"]
        else:
            ws.cell(row=r, column=1, value=label).font = sty["norm"]
            ws.cell(row=r, column=2, value=val).font = sty["bold"]
        r += 1
    sw(ws, [42, 15])

    # --- Security Findings ---
    ws2 = wb.create_sheet("Security Findings")
    hdr(ws2, 1, ["Priority", "Finding", "Evidence", "Impact", "Recommendation"], sty)
    r = 2
    security_items = [
        ("HIGH", "8 SRV rules allow 70 standard categories including weapons, extremism, streaming",
         "Rules: Artifactory, IMIS, Tableau, Service Mgmt, BigFix, SailPoint, GitLab, IBM AIX. Categories include WEAPONS_AND_BOMBS, MILITANCY_HATE_AND_EXTREMISM, STREAMING_MEDIA, REMOTE_ACCESS",
         "Servers allowed access to weapons, extremism, entertainment content. Appears copy-pasted from broad template.",
         "Trim category list to server-appropriate categories before consolidation"),
        ("HIGH", "Sharefile-Block-Uploads name says Block but action is ALLOW",
         "Rule name indicates DLP blocking of Sharefile uploads but action is ALLOW",
         "If this is a DLP control, it is not functioning. Users may believe uploads are blocked.",
         "Verify intent and fix name or action"),
        ("HIGH", f"{len(findings['broad_wildcards'])} broad wildcards in SSL bypass categories bypass entire cloud platforms",
         "Wildcards like .blob.core.windows.net, .amazonaws.com, .azure.com, .microsoft.com in SSL bypass categories",
         "Entire Azure, AWS, and Microsoft platforms bypass SSL inspection, creating massive blind spots",
         "Replace broad wildcards with specific subdomains where possible"),
        ("HIGH", f"OpenAI Azure endpoint in {len(findings['ai_urls'])} categories including SailPoint and DBA",
         ".npopenai-platform.openai.azure.com found in 8 categories: SSL bypass, SailPoint, DBA servers, NetSec, Microsoft Install/Channels",
         "AI platform URL silently added to server categories, bypassing SSL inspection and DLP controls",
         "Audit all categories containing this URL. Create dedicated AI/LLM category if AI access is approved."),
        ("HIGH", f"DLP Exceptions has {len(findings['dlp_broad'])} broad wildcard entries",
         f"Wildcards like .docusign.com, .irs.gov, .office.com, .sharefile.com, .transunion.com in DLP Exceptions (CUSTOM_13)",
         "DLP controls are waived for entire domains. 350 URLs total in DLP Exceptions.",
         "Audit DLP exceptions — replace broad wildcards with specific paths where possible"),
    ]
    # Medium items
    security_items += [
        ("MEDIUM", f"4 cross-policy category mismatches confirmed with URL evidence",
         "CUSTOM_88: Quickbase URLs used by Informatica rule. CUSTOM_247: PolyLens URLs used by Genesys bypass. CUSTOM_161: S3 URLs shared by Content Manager, Splunk, Data Center. CUSTOM_25: TransUnion URLs used by Anypoint bypass.",
         "Categories contain URLs for wrong services, creating unintended access or bypass paths",
         "Audit and split into dedicated per-service categories"),
        ("MEDIUM", f"{len(findings['content_mix'])} categories have mixed Microsoft/Azure content",
         "Gitlab-destinations has 57 Microsoft/Azure URLs. Splunk_Destinations has Office URLs. NetSec_JumpServers has docs.google.com.",
         "Service-specific categories are dumping grounds for unrelated URLs",
         "Extract Microsoft/Azure URLs into shared Microsoft category, keep service-specific URLs separate"),
        ("MEDIUM", f"{len(findings['shadows'])} rule shadowing instances",
         "Security Hacking Allow (order 22) shadows Standard User Category Block, Standard IoT Category Block, and Server Category Block with 9 overlapping categories",
         "ALLOW rule evaluated before BLOCK rules may be permitting traffic the BLOCK rules intend to deny",
         "Review rule ordering — consider moving BLOCK rules above shadowing ALLOW rules"),
        ("MEDIUM", "Custom URL category override behavior",
         "Blocked_Sites_IPs entries take precedence over Global_Allowed_Sites regardless of rule order. Specific URL matches do not consistently override wildcards.",
         "Creates confusion when troubleshooting. Allow rules may not work as expected.",
         "Document behavior. Check Blocked_Sites_IPs first when unblocking URLs."),
    ]
    # Low items
    security_items += [
        ("LOW", f"{len(findings['unused'])} unused categories not referenced by any rule",
         f"Categories not in any URL filtering or SSL inspection rule",
         "Clutter — may indicate orphaned categories from deleted rules",
         "Delete after confirming not pending deployment"),
        ("LOW", f"{len(findings['test_items'])} test/POC items still in production",
         "WildCard Testing, Content_Filter_TESTING, infra_testing, imaging_servers_test, GitLabPOC",
         "Test rules and categories should not persist in production",
         "Review and remove or rename"),
        ("LOW", f"E-sign tools scattered across blocked and allowed categories",
         "5 e-sign tools blocked (SignWell, PandaDoc, AssureSign, InsureSign, HelloSign). PandaDoc partially allowed. DocuSign in DLP exceptions. SignNow in Blocked_Sites_IPs.",
         "Inconsistent treatment — some blocked, some allowed, some in DLP exceptions",
         "Consolidate e-sign policy into clear allow/block decision"),
    ]

    for priority, finding, evidence, impact, rec in security_items:
        cel(ws2, r, 1, priority, sty, fill=L_RED if priority=="HIGH" else (L_ORANGE if priority=="MEDIUM" else L_YELLOW))
        cel(ws2, r, 2, finding, sty); cel(ws2, r, 3, evidence, sty)
        cel(ws2, r, 4, impact, sty); cel(ws2, r, 5, rec, sty)
        ws2.row_dimensions[r].height = 65; r += 1
    sw(ws2, [10, 42, 52, 48, 45]); ws2.freeze_panes = "A2"

    # --- Cross-Policy ---
    ws3 = wb.create_sheet("Cross-Policy Alignment")
    ws3["A1"] = f"Custom categories shared between URL filtering and SSL ({len(findings['shared_cats'])})"; ws3["A1"].font = sty["subtitle"]
    ws3.merge_cells("A1:G1")
    hdr(ws3, 2, ["Category", "Name", "URL Count", "Sample URLs", "URL Filtering Rules", "SSL Rules", "Issue?"], sty)
    r = 3
    for sc in findings["shared_cats"]:
        cel(ws3, r, 1, sc["id"], sty); cel(ws3, r, 2, sc["name"], sty); cel(ws3, r, 3, sc["url_count"], sty)
        cel(ws3, r, 4, trunc(sc["urls"], 5), sty); cel(ws3, r, 5, trunc(sc["url_rules"], 3), sty)
        cel(ws3, r, 6, trunc(sc["ssl_rules"], 3), sty)
        if sc["id"] in ("CUSTOM_88","CUSTOM_247","CUSTOM_161","CUSTOM_25"):
            cel(ws3, r, 7, "MISMATCH — different services", sty, fill=L_RED)
        else: cel(ws3, r, 7, "", sty)
        ws3.row_dimensions[r].height = 35; r += 1
    sw(ws3, [14, 25, 10, 45, 35, 35, 28]); ws3.freeze_panes = "A3"

    # --- Shadowing ---
    ws4 = wb.create_sheet("Rule Shadowing")
    hdr(ws4, 1, ["ALLOW Rule", "Order", "BLOCK Rule", "Order", "Overlapping Categories", "Count"], sty)
    r = 2
    for s in findings["shadows"]:
        cel(ws4, r, 1, s["allow"], sty); cel(ws4, r, 2, s["a_order"], sty)
        cel(ws4, r, 3, s["block"], sty); cel(ws4, r, 4, s["b_order"], sty)
        cel(ws4, r, 5, s["cats"], sty); cel(ws4, r, 6, s["count"], sty); r += 1
    sw(ws4, [35, 7, 35, 7, 55, 8]); ws4.freeze_panes = "A2"

    # --- Merge Opportunities ---
    ws5 = wb.create_sheet("Merge Opportunities")
    hdr(ws5, 1, ["Rules", "# Rules", "Location(s)", "Action"], sty)
    r = 2
    for m in sorted(findings["merges"], key=lambda x: -x["count"]):
        cel(ws5, r, 1, ", ".join(m["rules"]), sty); cel(ws5, r, 2, m["count"], sty)
        cel(ws5, r, 3, m["location"], sty); cel(ws5, r, 4, m["action"], sty); r += 1
    sw(ws5, [55, 8, 55, 10]); ws5.freeze_panes = "A2"

    # --- Category Content Issues ---
    ws6 = wb.create_sheet("Category Content Issues")
    ws6["A1"] = "Categories with content that doesn't match the category name"; ws6["A1"].font = sty["subtitle"]
    ws6.merge_cells("A1:E1")
    hdr(ws6, 2, ["Category", "Name", "Issue", "Microsoft/Azure URLs", "Total URLs"], sty)
    r = 3
    for cm in findings["content_mix"]:
        cel(ws6, r, 1, cm["id"], sty); cel(ws6, r, 2, cm["name"], sty)
        cel(ws6, r, 3, f"Contains {cm['ms_count']} Microsoft/Azure URLs in service-specific category", sty)
        cel(ws6, r, 4, cm["ms_count"], sty); cel(ws6, r, 5, cm["total"], sty); r += 1
    sw(ws6, [14, 30, 52, 18, 10]); ws6.freeze_panes = "A3"

    # --- Hit Count (if available) ---
    if url_hits:
        ws_h = wb.create_sheet("Policy Hit Count")
        ws_h["A1"] = "URL filtering rules — hit count from Web Insights logs"; ws_h["A1"].font = sty["subtitle"]
        ws_h.merge_cells("A1:D1")
        hdr(ws_h, 2, ["#", "Rule Name", "Hit Count", "% of Total"], sty)
        all_names = sorted(set(r["name"] for r in url_rules))
        total = sum(url_hits.values()) or 1
        wh = sorted([(n, url_hits.get(n, 0)) for n in all_names], key=lambda x: -x[1])
        r = 3
        for i, (name, count) in enumerate(wh, 1):
            cel(ws_h, r, 1, i, sty); cel(ws_h, r, 2, name, sty)
            cel(ws_h, r, 3, count, sty, fill=L_RED if count == 0 else None)
            cel(ws_h, r, 4, f"{count/total*100:.2f}%", sty); r += 1
        sw(ws_h, [5, 42, 12, 10]); ws_h.auto_filter.ref = f"A2:D{r-1}"; ws_h.freeze_panes = "A3"

    # --- Full Rule Inventory ---
    ws_inv = wb.create_sheet("Full Rule Inventory")
    hdr(ws_inv, 1, ["Order","Name","Action","State","Label","Std Cats","Custom Cats","Depts","Users","Locations","Description"], sty)
    r = 2
    for rd in sorted_rules:
        labels = ", ".join(l["name"] for l in rd.get("labels", []))
        std = [c for c in rd.get("urlCategories", []) if not c.startswith("CUSTOM_")]
        cust = [resolve(c, cat_map) for c in rd.get("urlCategories", []) if c.startswith("CUSTOM_")]
        depts = ", ".join(d["name"] for d in rd.get("departments", []))
        users = ", ".join(u["name"] for u in rd.get("users", []))
        locs = ", ".join(l["name"] for l in rd.get("locations", []))
        cel(ws_inv, r, 1, rd.get("order",""), sty); cel(ws_inv, r, 2, rd["name"], sty)
        cel(ws_inv, r, 3, rd.get("action",""), sty); cel(ws_inv, r, 4, rd.get("state",""), sty)
        cel(ws_inv, r, 5, labels, sty); cel(ws_inv, r, 6, len(std), sty)
        cel(ws_inv, r, 7, ", ".join(cust), sty); cel(ws_inv, r, 8, depts, sty)
        cel(ws_inv, r, 9, users, sty); cel(ws_inv, r, 10, locs, sty)
        cel(ws_inv, r, 11, rd.get("description",""), sty)
        if rd.get("state") == "DISABLED": ws_inv.cell(row=r, column=4).fill = L_RED
        r += 1
    ws_inv.auto_filter.ref = f"A1:K{r-1}"; ws_inv.freeze_panes = "C2"
    sw(ws_inv, [7, 38, 10, 10, 15, 8, 35, 28, 28, 45, 30])

    return wb


# === SSL INSPECTION EXCEL ===
def build_ssl_excel(ssl_rules, findings, cat_map, ssl_hits):
    wb = Workbook(); sty = SSL_STYLES
    sorted_rules = sorted(ssl_rules, key=lambda x: x.get("order", 999))

    # --- Summary ---
    ws = wb.active; ws.title = "Summary"
    ws["A1"] = "ZIA SSL Inspection Rules — Cleanup & Analysis"; ws["A1"].font = sty["title"]
    r = 3
    bp_true = sum(1 for x in ssl_rules if x["action"].get("doNotDecryptSubActions",{}).get("bypassOtherPolicies") == True)
    for label, val in [
        ("Total rules", len(ssl_rules)), ("DO_NOT_DECRYPT", sum(1 for x in ssl_rules if x["action"]["type"]=="DO_NOT_DECRYPT")),
        ("DECRYPT", sum(1 for x in ssl_rules if x["action"]["type"]=="DECRYPT")),
        ("Disabled", sum(1 for x in ssl_rules if x.get("state")=="DISABLED")),
        ("",""), ("Critical Findings",""),
        ("Non-human account inspection gap", "CONFIRMED"),
        ("bypassOtherPolicies=True rules", bp_true),
        ("Broad wildcards in bypass categories", len(findings["broad_wildcards"])),
        ("AI/LLM URLs in server categories", len(findings["ai_urls"])),
        ("DLP exception broad wildcards", len(findings["dlp_broad"])),
    ]:
        if val == "":
            ws.cell(row=r, column=1, value=label).font = sty["subtitle"] if label else sty["norm"]
        else:
            ws.cell(row=r, column=1, value=label).font = sty["norm"]
            c = ws.cell(row=r, column=2, value=val); c.font = sty["bold"]
            if label == "Non-human account inspection gap": c.font = sty["red"]
        r += 1
    sw(ws, [42, 15])

    # --- Security Findings ---
    ws2 = wb.create_sheet("Security Findings")
    hdr(ws2, 1, ["Priority", "Finding", "Evidence", "Impact", "Recommendation"], sty)
    r = 2
    items = [
        ("HIGH", "Rule 40 — non-human accounts bypass SSL inspection",
         "Default User SSL Inspect scoped to Zscaler-ZIA-Users AD group only. Testing confirmed accounts outside group bypass inspection. IAM cannot provision non-human accounts to this group.",
         "DLP, malware scanning, URL filtering do not apply to non-human traffic. All SSL-dependent controls ineffective.",
         "Option A: Catch-all DECRYPT at order 41. Option B: Remove group restriction. Option C (recommended): Flip default rule to DECRYPT."),
        ("HIGH", f"bypassOtherPolicies=True on {bp_true}/110 bypass rules",
         "102 DO_NOT_DECRYPT rules skip ALL Zscaler policies, not just SSL. Only 8 use bypassOtherPolicies=False.",
         "Traffic matching these rules has zero URL filtering, malware scanning, or DLP. Complete blind spots.",
         "Switch to False (Evaluate Other Policies) on per-user and location-only rules. Evaluate service-specific rules individually."),
        ("HIGH", f"{len(findings['broad_wildcards'])} broad wildcards bypass SSL for entire cloud platforms",
         ".blob.core.windows.net, .amazonaws.com, .azure.com, .microsoft.com, .google.com in bypass categories",
         f"Entire Azure, AWS, Microsoft, Google platforms bypass SSL inspection",
         "Replace with specific subdomains. Create shared Microsoft/Azure bypass category instead of duplicating."),
        ("HIGH", f"OpenAI Azure endpoint in {len(findings['ai_urls'])} categories",
         ".npopenai-platform.openai.azure.com in SailPoint, DBA servers, NetSec, SSL bypass, Microsoft Install/Channels",
         "AI platform URL in server categories bypasses SSL inspection and DLP. Data exfiltration risk.",
         "Create dedicated AI/LLM category. Remove from server-specific categories unless explicitly approved."),
        ("HIGH", f"DLP Exceptions has {len(findings['dlp_broad'])} broad wildcards",
         "Wildcards like .docusign.com, .irs.gov, .office.com, .sharefile.com in DLP Exceptions (350 URLs total)",
         "DLP controls waived for entire domains including government, financial, and file sharing services",
         "Audit — replace broad wildcards with specific paths"),
        ("MEDIUM", "SSL inspection expansion — partial coverage during batch rollout",
         "O365 Inspection rule being expanded via department batches (~9,300 users, Wednesday cadence). Not all users covered yet.",
         "Users not yet in AD group do not receive tenant restriction enforcement or O365 SSL inspection.",
         "Continue batch rollout. Exceptions via MyTech policy change form."),
        ("MEDIUM", "Tenant Restrictions V1 to V2 migration blocked",
         "V2 requires Entra cross-tenant access settings. Not yet configured. V1/V2 headers cannot coexist.",
         "Stuck on V1 until Identity team configures Entra. Clean cutover required.",
         "Coordinate with Identity team. Test on non-production tenant first."),
        ("MEDIUM", "IoT default is DO_NOT_DECRYPT",
         "Default SSL Inspection for IoT bypasses all IoT traffic.",
         "IoT devices are common attack vectors with zero traffic visibility.",
         "Evaluate inspection feasibility or add compensating controls."),
        ("MEDIUM", "Microsoft Form Bypass (order 1) has no scoping",
         "No URL categories, users, or locations. bypassOtherPolicies=True.",
         "First rule evaluated, matches broadly, complete blind spot for all MS Forms traffic.",
         "Add URL category scoping."),
    ]
    for priority, finding, evidence, impact, rec in items:
        cel(ws2, r, 1, priority, sty, fill=L_RED if priority=="HIGH" else L_ORANGE)
        cel(ws2, r, 2, finding, sty); cel(ws2, r, 3, evidence, sty)
        cel(ws2, r, 4, impact, sty); cel(ws2, r, 5, rec, sty)
        ws2.row_dimensions[r].height = 65; r += 1
    sw(ws2, [10, 42, 52, 48, 48]); ws2.freeze_panes = "A2"

    # --- DECRYPT Coverage ---
    ws3 = wb.create_sheet("DECRYPT Coverage")
    hdr(ws3, 1, ["Rule", "Order", "Labels", "Groups", "Predefined"], sty)
    r = 2
    for dc in findings["decrypt_rules"]:
        cel(ws3, r, 1, dc["name"], sty); cel(ws3, r, 2, dc["order"], sty)
        cel(ws3, r, 3, dc["labels"], sty); cel(ws3, r, 4, dc["groups"], sty)
        cel(ws3, r, 5, "Yes" if dc["predefined"] else "", sty); r += 1
    r += 1
    ws3.cell(row=r, column=1, value="Label Breakdown").font = sty["subtitle"]; r += 1
    hdr(ws3, r, ["Label", "DECRYPT", "BYPASS", "Skip All Policies", "Evaluate Other"], sty); r += 1
    for label in sorted(findings["label_stats"].keys()):
        s = findings["label_stats"][label]
        cel(ws3, r, 1, label, sty); cel(ws3, r, 2, s["decrypt"], sty)
        cel(ws3, r, 3, s["bypass"], sty); cel(ws3, r, 4, s["bypass_all"], sty)
        cel(ws3, r, 5, s["bypass_eval"], sty); r += 1
    sw(ws3, [35, 8, 20, 18, 16])

    # --- Broad Wildcards ---
    ws4 = wb.create_sheet("Broad Wildcards in Bypass")
    hdr(ws4, 1, ["Wildcard", "Category", "Category ID"], sty)
    r = 2
    for bw in sorted(findings["broad_wildcards"], key=lambda x: x["url"]):
        cel(ws4, r, 1, bw["url"], sty); cel(ws4, r, 2, bw["cat_name"], sty)
        cel(ws4, r, 3, bw["cat_id"], sty); r += 1
    sw(ws4, [30, 40, 14]); ws4.freeze_panes = "A2"

    # --- Hit Count ---
    if ssl_hits:
        ws_h = wb.create_sheet("Policy Hit Count")
        hdr(ws_h, 1, ["#", "Rule Name", "Hit Count", "% of Total"], sty)
        all_names = sorted(set(r["name"] for r in ssl_rules))
        total = sum(ssl_hits.values()) or 1
        wh = sorted([(n, ssl_hits.get(n, 0)) for n in all_names], key=lambda x: -x[1])
        r = 2
        for i, (name, count) in enumerate(wh, 1):
            cel(ws_h, r, 1, i, sty); cel(ws_h, r, 2, name, sty)
            cel(ws_h, r, 3, count, sty, fill=L_RED if count == 0 else None)
            cel(ws_h, r, 4, f"{count/total*100:.2f}%", sty); r += 1
        sw(ws_h, [5, 42, 12, 10]); ws_h.auto_filter.ref = f"A1:D{r-1}"; ws_h.freeze_panes = "A2"

    # --- Full Inventory ---
    ws_inv = wb.create_sheet("Full Rule Inventory")
    hdr(ws_inv, 1, ["Order","Name","Action","State","Predefined","Label","bypassOtherPolicies","Std Cats","Custom Cats","Users","Groups","Locations"], sty)
    r = 2
    for rd in sorted_rules:
        act = rd["action"]
        labels = ", ".join(l["name"] for l in rd.get("labels",[]))
        std = ", ".join(c for c in rd.get("urlCategories",[]) if not c.startswith("CUSTOM_"))
        cust = ", ".join(resolve(c, cat_map) for c in rd.get("urlCategories",[]) if c.startswith("CUSTOM_"))
        users = ", ".join(u["name"] for u in rd.get("users",[]))
        groups = ", ".join(g["name"] for g in rd.get("groups",[]))
        locs = ", ".join(l["name"] for l in rd.get("locations",[]))
        bp = act.get("doNotDecryptSubActions",{}).get("bypassOtherPolicies","")
        cel(ws_inv, r, 1, rd.get("order",""), sty); cel(ws_inv, r, 2, rd["name"], sty)
        cel(ws_inv, r, 3, act["type"], sty); cel(ws_inv, r, 4, rd.get("state",""), sty)
        cel(ws_inv, r, 5, "Yes" if rd.get("predefined") else "", sty)
        cel(ws_inv, r, 6, labels, sty); cel(ws_inv, r, 7, str(bp) if bp != "" else "", sty)
        cel(ws_inv, r, 8, std, sty); cel(ws_inv, r, 9, cust, sty)
        cel(ws_inv, r, 10, users, sty); cel(ws_inv, r, 11, groups, sty); cel(ws_inv, r, 12, locs, sty)
        if rd.get("state") == "DISABLED": ws_inv.cell(row=r, column=4).fill = L_RED
        r += 1
    ws_inv.auto_filter.ref = f"A1:L{r-1}"; ws_inv.freeze_panes = "C2"
    sw(ws_inv, [7, 36, 18, 10, 10, 15, 17, 25, 35, 35, 20, 50])

    return wb


# === WORD DOCS ===
def build_url_docx(findings):
    doc = DocxDoc()
    s = doc.styles["Normal"]; s.font.name = "Calibri"; s.font.size = Pt(11)
    doc.add_heading("Zscaler URL Filtering Rules", level=0)
    doc.add_heading("Cleanup & Consolidation Proposal", level=1)
    doc.add_paragraph("Network Security Engineering — Security Control Management")

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        f"Analysis of {findings.get('url_rule_count', 151)} URL filtering rules with cross-referencing "
        f"against {findings.get('cat_count', 261)} custom URL category contents identified critical "
        f"security findings, policy misconfigurations, and consolidation opportunities."
    )

    doc.add_heading("Critical Security Findings", level=1)

    doc.add_heading("Broad SRV Category Allowlists", level=2)
    doc.add_paragraph(
        "Eight server rules allow 70 standard URL categories including WEAPONS_AND_BOMBS, "
        "MILITANCY_HATE_AND_EXTREMISM, STREAMING_MEDIA, and REMOTE_ACCESS. Servers have no "
        "legitimate need for these categories. Lists should be trimmed before consolidation."
    )

    doc.add_heading("Broad Wildcards in SSL Bypass Categories", level=2)
    doc.add_paragraph(
        f"{len(findings['broad_wildcards'])} broad wildcards like .blob.core.windows.net, .amazonaws.com, "
        f".azure.com, and .microsoft.com exist in SSL bypass categories. These bypass SSL inspection "
        f"for entire cloud platforms, creating massive blind spots for DLP, malware, and URL filtering."
    )

    doc.add_heading("OpenAI Azure Endpoint in Server Categories", level=2)
    doc.add_paragraph(
        f"The URL .npopenai-platform.openai.azure.com appears in {len(findings['ai_urls'])} categories "
        f"including SailPoint, DBA servers, and NetSec JumpServers. This AI platform URL was likely "
        f"added as part of a Microsoft URL bundle and now bypasses SSL inspection and DLP controls "
        f"through server-specific categories."
    )

    doc.add_heading("DLP Exceptions Audit", level=2)
    doc.add_paragraph(
        f"The DLP Exceptions category contains 350 URLs with {len(findings['dlp_broad'])} broad "
        f"wildcards including .docusign.com, .irs.gov, .office.com, and .sharefile.com. DLP controls "
        f"are waived for these entire domains."
    )

    doc.add_heading("Cross-Policy Category Mismatches", level=2)
    doc.add_paragraph(
        "4 custom URL categories are shared between URL filtering and SSL inspection but contain "
        "URLs for different services than the rules referencing them suggest. CUSTOM_88 contains "
        "Quickbase URLs but is used by an Informatica rule. CUSTOM_247 contains PolyLens URLs but "
        "is used by a Genesys bypass."
    )

    doc.add_heading("Category Content Analysis", level=1)

    doc.add_heading("Massive URL Duplication", level=2)
    doc.add_paragraph(
        f"{len(findings['duplicates'])} URLs appear in multiple categories. graph.microsoft.com appears "
        f"in 32 categories. This means updating a Microsoft URL requires touching 20+ categories."
    )

    for cm in findings["content_mix"]:
        doc.add_paragraph(
            f'{cm["name"]} ({cm["id"]}): {cm["ms_count"]} Microsoft/Azure URLs in a '
            f'{cm["total"]}-URL service-specific category.',
            style="List Bullet"
        )

    doc.add_heading("Cleanup Plan", level=1)
    doc.add_paragraph("Phase 1: Fix security findings (Sharefile, broad SRV categories, AI URLs, DLP exceptions). Phase 2: Merge duplicate-scoped rules. Phase 3: Consolidate per-server SRV rules. Phase 4: Address category hygiene (duplicates, unused, test items).")

    doc.add_heading("Recommendation", level=1)
    doc.add_paragraph(
        "Address security findings first, then proceed through phased cleanup. Establish a shared "
        "Microsoft/Azure URL category to eliminate cross-category duplication. Create a dedicated "
        "AI/LLM category. Audit DLP Exceptions and Blocked_Sites for stale entries. Document the "
        "custom URL category override behavior."
    )
    return doc

def build_ssl_docx(findings):
    doc = DocxDoc()
    s = doc.styles["Normal"]; s.font.name = "Arial"; s.font.size = Pt(10)
    doc.add_heading("Zscaler SSL Inspection Rules", level=0)
    doc.add_heading("Cleanup & Security Posture Proposal", level=1)
    doc.add_paragraph("Network Security Engineering — Security Control Management")

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        f"Analysis of {findings.get('ssl_rule_count', 114)} SSL inspection rules identified a confirmed "
        f"critical security gap, an inverted inspection posture, and significant blind spots from broad "
        f"wildcards and bypassOtherPolicies overuse."
    )

    doc.add_heading("Critical Finding — Non-Human Account Gap", level=1)
    doc.add_paragraph(
        "CONFIRMED: Rule 40 (Default User SSL Inspect) is scoped to Zscaler-ZIA-Users AD group only. "
        "Accounts not in this group bypass SSL inspection entirely. IAM confirmed non-human accounts "
        "cannot be provisioned through Okta sync."
    )
    doc.add_paragraph("Fix options: (A) Catch-all DECRYPT at order 41, (B) Remove group restriction, (C) Flip default to DECRYPT (recommended).")

    doc.add_heading("Broad Wildcards in Bypass Categories", level=1)
    doc.add_paragraph(
        f"{len(findings['broad_wildcards'])} broad wildcards bypass SSL for entire cloud platforms. "
        f"Entries like .blob.core.windows.net, .amazonaws.com, .azure.com, and .microsoft.com exist "
        f"in SSL bypass categories, creating massive uninspected traffic paths."
    )

    doc.add_heading("bypassOtherPolicies Overuse", level=1)
    bp_true = sum(1 for r in findings.get("ssl_rules_raw", [])
                  if r["action"].get("doNotDecryptSubActions",{}).get("bypassOtherPolicies") == True)
    doc.add_paragraph(
        f"{bp_true} of 110 bypass rules skip ALL Zscaler policies. Per-user and location-only bypasses "
        f"should switch to bypassOtherPolicies=False to maintain URL filtering and DLP."
    )

    doc.add_heading("Cleanup Plan", level=1)
    doc.add_paragraph("Phase 1: Naming fixes (13 items). Phase 2: Consolidate 19 per-user bypasses. Phase 3: Merge location-only and GitLab rules. Strategic: Flip default to DECRYPT, address bypassOtherPolicies.")

    doc.add_heading("Recommendation", level=1)
    doc.add_paragraph(
        "Close the non-human account gap (Option C). Switch per-user bypass rules to bypassOtherPolicies=False. "
        "Replace broad wildcards with specific subdomains. Strategic goal: inspect-by-default posture "
        "with documented exceptions."
    )
    return doc


# === MAIN ===
def main():
    if len(sys.argv) < 3:
        print(__doc__); sys.exit(1)

    url_path, ssl_path = sys.argv[1], sys.argv[2]
    folder = os.path.dirname(os.path.abspath(url_path)) or os.getcwd()

    for p in [url_path, ssl_path]:
        try: open(p, "r").close(); print(f"  OK: {p}")
        except Exception as e: print(f"ERROR: {p}: {e}"); sys.exit(1)

    print("\nStep 1: Loading data...")
    with open(url_path, "r", encoding="utf-8") as f: url_rules = json.load(f)
    with open(ssl_path, "r", encoding="utf-8") as f: ssl_rules = json.load(f)
    print(f"  URL filtering: {len(url_rules)} rules")
    print(f"  SSL inspection: {len(ssl_rules)} rules")

    cat_map, cat_list = load_categories(folder)
    if not cat_map:
        print("ERROR: url_cat.json not found. Place it in the same folder.")
        sys.exit(1)
    print(f"  Custom categories: {len(cat_map)}")

    print("\nStep 2: Loading Web Insights logs...")
    url_hits, ssl_hits = load_web_logs(folder)
    if not url_hits and not ssl_hits:
        print("  No CSV logs found (optional).")

    print("\nStep 3: Running analysis...")
    findings = analyze_all(url_rules, ssl_rules, cat_map, cat_list)
    findings["url_rule_count"] = len(url_rules)
    findings["ssl_rule_count"] = len(ssl_rules)
    findings["cat_count"] = len(cat_map)
    findings["ssl_rules_raw"] = ssl_rules

    print(f"  Shadows: {len(findings['shadows'])}")
    print(f"  Merges: {len(findings['merges'])}")
    print(f"  Broad wildcards: {len(findings['broad_wildcards'])}")
    print(f"  AI URLs: {len(findings['ai_urls'])}")
    print(f"  DLP broad wildcards: {len(findings['dlp_broad'])}")
    print(f"  Duplicate URLs: {len(findings['duplicates'])}")
    print(f"  Content mismatches: {len(findings['content_mix'])}")

    d = os.getcwd()

    print("\nStep 4: Building URL Filtering Excel...")
    wb1 = build_url_excel(url_rules, findings, cat_map, url_hits)
    p1 = os.path.join(d, "ZIA_URL_Filtering_Cleanup_Strategy.xlsx")
    wb1.save(p1); print(f"  Saved: {p1}")

    print("\nStep 5: Building SSL Inspection Excel...")
    wb2 = build_ssl_excel(ssl_rules, findings, cat_map, ssl_hits)
    p2 = os.path.join(d, "ZIA_SSL_Inspection_Cleanup_Strategy.xlsx")
    wb2.save(p2); print(f"  Saved: {p2}")

    if HAS_DOCX:
        print("\nStep 6: Building Word docs...")
        doc1 = build_url_docx(findings)
        dp1 = os.path.join(d, "ZIA_URL_Filtering_Proposal.docx")
        doc1.save(dp1); print(f"  Saved: {dp1}")

        doc2 = build_ssl_docx(findings)
        dp2 = os.path.join(d, "ZIA_SSL_Inspection_Proposal.docx")
        doc2.save(dp2); print(f"  Saved: {dp2}")

    print("\nDone! Generated files:")
    print("  - ZIA_URL_Filtering_Cleanup_Strategy.xlsx")
    print("  - ZIA_SSL_Inspection_Cleanup_Strategy.xlsx")
    if HAS_DOCX:
        print("  - ZIA_URL_Filtering_Proposal.docx")
        print("  - ZIA_SSL_Inspection_Proposal.docx")

if __name__ == "__main__":
    main()
