"""
ZIA URL Filtering & SSL Inspection — Deep Analysis
------------------------------------------------------
Reads policy JSONs, custom category JSON, and optionally Web Insights
CSV logs for hit count analysis. Shows policies with 0 hits.

PowerShell:
    pip install openpyxl python-docx
    python zia_policy_deep_analysis.py url_filtering.json sslpol.json

With hit count logs (optional, auto-detected from folder):
    Just place any *_WEB_log.csv files in the same folder.
    The script will find and process them automatically.

With custom categories (optional):
    Place url_cat.json in the same folder for category name resolution.

Output (in current directory):
    - ZIA_Policy_Deep_Analysis.xlsx
    - ZIA_Policy_Deep_Analysis.docx
"""

import json, sys, os, re, glob
from datetime import datetime
from collections import Counter, defaultdict

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
except ImportError:
    print("ERROR: openpyxl required. Install with: pip install openpyxl")
    sys.exit(1)

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("NOTE: python-docx not installed — Word doc will be skipped.")

# Styles
BOLD = Font(name="Calibri", bold=True, size=11)
NORMAL = Font(name="Calibri", size=11)
TITLE = Font(name="Calibri", bold=True, size=14)
SUBTITLE = Font(name="Calibri", bold=True, size=12)
HEADER_FILL = PatternFill("solid", fgColor="404040")
HEADER_FONT = Font(name="Calibri", bold=True, size=11, color="FFFFFF")
LIGHT_RED = PatternFill("solid", fgColor="F2DCDB")
LIGHT_ORANGE = PatternFill("solid", fgColor="FDE9D9")
LIGHT_YELLOW = PatternFill("solid", fgColor="FFFFCC")
LIGHT_GREEN = PatternFill("solid", fgColor="D8E4BC")
BDR = Border(left=Side(style="thin", color="AAAAAA"), right=Side(style="thin", color="AAAAAA"),
             top=Side(style="thin", color="AAAAAA"), bottom=Side(style="thin", color="AAAAAA"))

def hdr(ws, row, headers):
    for c, h in enumerate(headers, 1):
        cl = ws.cell(row=row, column=c, value=h)
        cl.font = HEADER_FONT; cl.fill = HEADER_FILL; cl.border = BDR
        cl.alignment = Alignment(wrap_text=True, vertical="center")

def cel(ws, row, col, val, font=None, fill=None):
    c = ws.cell(row=row, column=col, value=val)
    c.font = font or NORMAL; c.border = BDR
    if fill: c.fill = fill
    c.alignment = Alignment(wrap_text=True, vertical="top")
    return c

def set_widths(ws, widths):
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w


def load_cat_map(folder):
    """Load custom category name mapping from url_cat.json if present."""
    cat_path = os.path.join(folder, "url_cat.json")
    if not os.path.exists(cat_path):
        # Try alternate names
        for name in ["zia_custom_categories.json", "custom_categories.json", "categories.json"]:
            alt = os.path.join(folder, name)
            if os.path.exists(alt):
                cat_path = alt
                break
        else:
            return {}

    print(f"  Loading category names from {os.path.basename(cat_path)}...")
    with open(cat_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    cat_map = {}
    cats = data.get("categories", data) if isinstance(data, dict) else data
    if isinstance(cats, dict):
        cats = list(cats.values()) if not any(isinstance(v, str) for v in cats.values()) else []
    for c in cats:
        if isinstance(c, dict) and "id" in c:
            cat_map[c["id"]] = c.get("name", c.get("configuredName", c["id"]))
    print(f"  Loaded {len(cat_map)} category names")
    return cat_map


def resolve_cat(cat_id, cat_map):
    if cat_id.startswith("CUSTOM_") and cat_id in cat_map:
        return f"{cat_map[cat_id]} ({cat_id})"
    return cat_id


def load_web_logs(folder):
    """Auto-detect and load Web Insights CSV logs from folder."""
    csv_files = glob.glob(os.path.join(folder, "*WEB_log*")) + \
                glob.glob(os.path.join(folder, "*web_log*")) + \
                glob.glob(os.path.join(folder, "SSL_inspection*")) + \
                glob.glob(os.path.join(folder, "*.csv"))

    # Deduplicate
    csv_files = list(set(csv_files))
    if not csv_files:
        return {}, {}

    print(f"  Found {len(csv_files)} CSV file(s)")

    url_hits = Counter()  # URL filtering policy name -> count
    ssl_hits = Counter()  # SSL policy name -> count

    for csv_file in csv_files:
        try:
            # Detect encoding
            with open(csv_file, "rb") as f:
                raw = f.read(500)
            if raw[:2] == b"\xff\xfe": enc = "utf-16-le"
            elif raw[:2] == b"\xfe\xff": enc = "utf-16-be"
            elif raw[:3] == b"\xef\xbb\xbf": enc = "utf-8-sig"
            else: enc = "utf-8"

            with open(csv_file, "r", encoding=enc, errors="replace") as f:
                lines = f.readlines()

            # Find header row
            header_idx = 0
            for i, line in enumerate(lines):
                lower = line.lower()
                if any(kw in lower for kw in ["event time", "policy action", "url filtering policy", "ssl policy"]):
                    header_idx = i
                    break
                if i > 15: break

            if header_idx >= len(lines): continue

            # Parse header
            sep = "\t" if lines[header_idx].count("\t") > lines[header_idx].count(",") else ","
            headers = [h.strip().strip('"') for h in lines[header_idx].split(sep)]

            # Find policy columns
            url_pol_col = None
            ssl_pol_col = None
            for idx, h in enumerate(headers):
                hl = h.lower()
                if "url filtering policy" in hl or "url filtering rule" in hl:
                    url_pol_col = idx
                elif "ssl policy" in hl or "ssl inspection policy" in hl:
                    ssl_pol_col = idx

            if url_pol_col is None and ssl_pol_col is None:
                print(f"    {os.path.basename(csv_file)}: no policy name column found, skipping")
                continue

            # Count hits
            file_url_hits = 0
            file_ssl_hits = 0
            for line in lines[header_idx + 1:]:
                fields = [f.strip().strip('"') for f in line.split(sep)]
                if url_pol_col is not None and url_pol_col < len(fields):
                    val = fields[url_pol_col]
                    if val and val.lower() not in ("none", "", "na"):
                        url_hits[val] += 1
                        file_url_hits += 1
                if ssl_pol_col is not None and ssl_pol_col < len(fields):
                    val = fields[ssl_pol_col]
                    if val and val.lower() not in ("none", "", "na"):
                        ssl_hits[val] += 1
                        file_ssl_hits += 1

            parts = []
            if file_url_hits: parts.append(f"{file_url_hits:,} URL hits")
            if file_ssl_hits: parts.append(f"{file_ssl_hits:,} SSL hits")
            print(f"    {os.path.basename(csv_file)}: {', '.join(parts) if parts else 'no policy hits found'}")

        except Exception as e:
            print(f"    {os.path.basename(csv_file)}: ERROR {e}")

    return dict(url_hits), dict(ssl_hits)


def analyze_url_rules(url_rules, cat_map):
    sorted_rules = sorted(url_rules, key=lambda x: x.get("order", 999))
    findings = {}

    # Rule shadowing
    shadows = []
    blocks = [(r["order"], r) for r in sorted_rules if r.get("action") == "BLOCK"]
    allows = [(r["order"], r) for r in sorted_rules if r.get("action") == "ALLOW"]
    for b_ord, b_r in blocks:
        b_cats = set(b_r.get("urlCategories", []))
        if not b_cats: continue
        for a_ord, a_r in allows:
            if a_ord < b_ord:
                overlap = set(a_r.get("urlCategories", [])) & b_cats
                if overlap and len(overlap) >= 3:
                    shadows.append({
                        "allow_name": a_r["name"], "allow_order": a_ord,
                        "block_name": b_r["name"], "block_order": b_ord,
                        "overlap_count": len(overlap),
                        "overlap_cats": ", ".join(resolve_cat(c, cat_map) for c in sorted(overlap)[:8]),
                    })
    findings["shadows"] = shadows

    # Duplicate scoping
    scope_groups = defaultdict(list)
    for r in sorted_rules:
        locs = tuple(sorted(l["name"] for l in r.get("locations", [])))
        if locs:
            depts = tuple(sorted(d["name"] for d in r.get("departments", [])))
            users = tuple(sorted(u["name"] for u in r.get("users", [])))
            key = (locs, depts, users, r.get("action", ""))
            scope_groups[key].append(r)
    findings["merges"] = [
        {"rules": [r["name"] for r in rules], "location": ", ".join(key[0])[:80], "action": key[3], "count": len(rules)}
        for key, rules in scope_groups.items() if len(rules) > 1
    ]

    # Admin activity
    admin_counts = Counter(r.get("lastModifiedBy", {}).get("name", "unknown") for r in url_rules)
    findings["admins"] = admin_counts.most_common(10)

    # Timeline
    timeline = Counter()
    oldest, newest = None, None
    for r in url_rules:
        ts = r.get("lastModifiedTime", 0)
        if ts:
            dt = datetime.fromtimestamp(ts)
            timeline[dt.strftime("%Y-%m")] += 1
            if oldest is None or dt < oldest: oldest = dt
            if newest is None or dt > newest: newest = dt
    findings["timeline"] = {"months": dict(sorted(timeline.items())), "oldest": oldest, "newest": newest}

    # Location distribution
    loc_counts = Counter()
    for r in url_rules:
        for l in r.get("locations", []):
            loc_counts[l["name"]] += 1
    findings["locations"] = {"counts": loc_counts.most_common(20), "unique": len(loc_counts),
                             "no_location": sum(1 for r in url_rules if not r.get("locations"))}

    # Category frequency
    cat_counts = Counter()
    for r in url_rules:
        for c in r.get("urlCategories", []):
            cat_counts[c] += 1
    findings["categories"] = {"most_used": cat_counts.most_common(10),
                              "single_use_count": sum(1 for c, n in cat_counts.items() if n == 1)}

    # Broad rules
    findings["broad_rules"] = [
        {"name": r["name"], "order": r.get("order"), "action": r.get("action")}
        for r in sorted_rules
        if not r.get("urlCategories") and not r.get("users") and not r.get("locations") and not r.get("departments")
    ]

    return findings


def analyze_ssl_rules(ssl_rules):
    sorted_rules = sorted(ssl_rules, key=lambda x: x.get("order", 999))
    findings = {}

    # DECRYPT coverage
    findings["decrypt_coverage"] = [
        {"name": r["name"], "order": r.get("order"),
         "labels": ", ".join(l["name"] for l in r.get("labels", [])) or "(none)",
         "groups": ", ".join(g["name"] for g in r.get("groups", [])) or "(none)",
         "predefined": r.get("predefined", False)}
        for r in sorted_rules if r["action"]["type"] == "DECRYPT"
    ]

    # Label stats
    label_stats = defaultdict(lambda: {"decrypt": 0, "bypass": 0, "bypass_all": 0, "bypass_eval": 0})
    for r in sorted_rules:
        for l in r.get("labels", []):
            if r["action"]["type"] == "DECRYPT":
                label_stats[l["name"]]["decrypt"] += 1
            else:
                label_stats[l["name"]]["bypass"] += 1
                bp = r["action"].get("doNotDecryptSubActions", {}).get("bypassOtherPolicies", "")
                if bp == True: label_stats[l["name"]]["bypass_all"] += 1
                elif bp == False: label_stats[l["name"]]["bypass_eval"] += 1
    findings["label_stats"] = dict(label_stats)
    findings["no_label_count"] = sum(1 for r in sorted_rules if not r.get("labels"))

    # Location overlap
    loc_rules = defaultdict(list)
    for r in sorted_rules:
        for l in r.get("locations", []):
            loc_rules[l["name"]].append({"name": r["name"], "action": r["action"]["type"]})
    findings["location_overlap"] = {l: rs for l, rs in loc_rules.items() if len(rs) > 1}

    # Admins
    findings["admins"] = Counter(r.get("lastModifiedBy", {}).get("name", "unknown") for r in ssl_rules).most_common(10)

    # Timeline
    tl = Counter()
    for r in ssl_rules:
        ts = r.get("lastModifiedTime", 0)
        if ts: tl[datetime.fromtimestamp(ts).strftime("%Y-%m")] += 1
    findings["timeline"] = dict(sorted(tl.items()))

    return findings


def build_excel(url_findings, ssl_findings, url_rules, ssl_rules, url_hits, ssl_hits, cat_map):
    wb = Workbook()

    # --- TAB 1: Summary ---
    ws1 = wb.active; ws1.title = "Analysis Summary"
    ws1["A1"] = "ZIA URL Filtering & SSL Inspection — Deep Analysis"
    ws1["A1"].font = TITLE

    r = 3
    ws1.cell(row=r, column=1, value="URL Filtering").font = SUBTITLE; r += 1
    for label, val in [
        ("Rule shadowing candidates", len(url_findings["shadows"])),
        ("Duplicate-scope merge opportunities", len(url_findings["merges"])),
        ("Broadly scoped rules", len(url_findings["broad_rules"])),
        ("Categories used in only 1 rule", url_findings["categories"]["single_use_count"]),
        ("Unique locations", url_findings["locations"]["unique"]),
        ("Rules with no location", url_findings["locations"]["no_location"]),
    ]:
        ws1.cell(row=r, column=1, value=label).font = NORMAL
        ws1.cell(row=r, column=2, value=val).font = BOLD; r += 1

    if url_hits:
        r += 1
        ws1.cell(row=r, column=1, value="URL Filtering Hit Count").font = SUBTITLE; r += 1
        total_url_rules = len(url_rules)
        rules_with_hits = sum(1 for rn in set(r["name"] for r in url_rules) if rn in url_hits)
        ws1.cell(row=r, column=1, value="Total rules").font = NORMAL
        ws1.cell(row=r, column=2, value=total_url_rules).font = BOLD; r += 1
        ws1.cell(row=r, column=1, value="Rules with hits in logs").font = NORMAL
        ws1.cell(row=r, column=2, value=rules_with_hits).font = BOLD; r += 1
        ws1.cell(row=r, column=1, value="Rules with 0 hits").font = NORMAL
        ws1.cell(row=r, column=2, value=total_url_rules - rules_with_hits).font = BOLD; r += 1

    r += 1
    ws1.cell(row=r, column=1, value="SSL Inspection").font = SUBTITLE; r += 1
    for label, val in [
        ("DECRYPT rules", len(ssl_findings["decrypt_coverage"])),
        ("Locations with overlapping bypasses", len(ssl_findings["location_overlap"])),
        ("Rules with no label", ssl_findings["no_label_count"]),
    ]:
        ws1.cell(row=r, column=1, value=label).font = NORMAL
        ws1.cell(row=r, column=2, value=val).font = BOLD; r += 1

    for label in sorted(ssl_findings["label_stats"].keys()):
        s = ssl_findings["label_stats"][label]
        ws1.cell(row=r, column=1, value=f"  {label}: DECRYPT / BYPASS").font = NORMAL
        ws1.cell(row=r, column=2, value=f'{s["decrypt"]} / {s["bypass"]}').font = BOLD; r += 1

    if ssl_hits:
        r += 1
        ws1.cell(row=r, column=1, value="SSL Inspection Hit Count").font = SUBTITLE; r += 1
        total_ssl = len(ssl_rules)
        ssl_with_hits = sum(1 for rn in set(r["name"] for r in ssl_rules) if rn in ssl_hits)
        ws1.cell(row=r, column=1, value="Total rules").font = NORMAL
        ws1.cell(row=r, column=2, value=total_ssl).font = BOLD; r += 1
        ws1.cell(row=r, column=1, value="Rules with hits").font = NORMAL
        ws1.cell(row=r, column=2, value=ssl_with_hits).font = BOLD; r += 1
        ws1.cell(row=r, column=1, value="Rules with 0 hits").font = NORMAL
        ws1.cell(row=r, column=2, value=total_ssl - ssl_with_hits).font = BOLD; r += 1

    set_widths(ws1, [45, 20])

    # --- TAB 2: Rule Shadowing ---
    ws2 = wb.create_sheet("Rule Shadowing")
    ws2["A1"] = "ALLOW rules evaluated before BLOCK rules with overlapping categories"
    ws2["A1"].font = SUBTITLE; ws2.merge_cells("A1:F1")
    hdr(ws2, 2, ["ALLOW Rule", "Order", "BLOCK Rule", "Order", "Overlapping Categories", "Count"])
    r = 3
    for s in url_findings["shadows"]:
        cel(ws2, r, 1, s["allow_name"]); cel(ws2, r, 2, s["allow_order"])
        cel(ws2, r, 3, s["block_name"]); cel(ws2, r, 4, s["block_order"])
        cel(ws2, r, 5, s["overlap_cats"]); cel(ws2, r, 6, s["overlap_count"])
        r += 1
    set_widths(ws2, [35, 7, 35, 7, 55, 8]); ws2.freeze_panes = "A3"

    # --- TAB 3: Merge Opportunities ---
    ws3 = wb.create_sheet("Merge Opportunities")
    ws3["A1"] = "Rules with identical scoping (same locations, users, departments, action)"
    ws3["A1"].font = SUBTITLE; ws3.merge_cells("A1:D1")
    hdr(ws3, 2, ["Rules (same scope)", "# Rules", "Location(s)", "Action"])
    r = 3
    for m in sorted(url_findings["merges"], key=lambda x: -x["count"]):
        cel(ws3, r, 1, ", ".join(m["rules"])); cel(ws3, r, 2, m["count"])
        cel(ws3, r, 3, m["location"]); cel(ws3, r, 4, m["action"])
        ws3.row_dimensions[r].height = 30; r += 1
    set_widths(ws3, [55, 8, 55, 10]); ws3.freeze_panes = "A3"

    # --- TAB 4: URL Policy Hit Count ---
    if url_hits or ssl_hits:
        ws4 = wb.create_sheet("URL Policy Hit Count")
        ws4["A1"] = "URL filtering rules — hit count from Web Insights logs"
        ws4["A1"].font = SUBTITLE; ws4.merge_cells("A1:D1")
        ws4["A2"] = "Rules with 0 hits may be unused or only triggered outside the log window"
        ws4["A2"].font = Font(name="Calibri", size=10, italic=True); ws4.merge_cells("A2:D2")
        hdr(ws4, 3, ["#", "Rule Name", "Hit Count", "% of Total"])

        all_url_names = sorted(set(r["name"] for r in url_rules))
        total = sum(url_hits.values()) or 1
        r = 4
        # Sort: hits descending, then 0-hit rules
        with_hits = [(name, url_hits.get(name, 0)) for name in all_url_names]
        with_hits.sort(key=lambda x: -x[1])

        for i, (name, count) in enumerate(with_hits, 1):
            cel(ws4, r, 1, i)
            cel(ws4, r, 2, name)
            cel(ws4, r, 3, count)
            pct = count / total * 100
            cel(ws4, r, 4, f"{pct:.2f}%")
            if count == 0:
                ws4.cell(row=r, column=3).fill = LIGHT_RED
            r += 1

        # Also add rules from logs that don't match any policy (possible renames)
        log_only = set(url_hits.keys()) - set(all_url_names)
        if log_only:
            r += 1
            ws4.cell(row=r, column=1, value="Rules in logs but not in current policy (possibly renamed/deleted)").font = BOLD
            ws4.merge_cells(f"A{r}:D{r}"); r += 1
            for name in sorted(log_only):
                cel(ws4, r, 2, name, fill=LIGHT_YELLOW)
                cel(ws4, r, 3, url_hits[name]); r += 1

        set_widths(ws4, [5, 42, 12, 10])
        ws4.auto_filter.ref = f"A3:D{r-1}"; ws4.freeze_panes = "A4"

    # --- TAB 5: SSL Policy Hit Count ---
    if ssl_hits:
        ws5 = wb.create_sheet("SSL Policy Hit Count")
        ws5["A1"] = "SSL inspection rules — hit count from Web Insights logs"
        ws5["A1"].font = SUBTITLE; ws5.merge_cells("A1:D1")
        ws5["A2"] = "Rules with 0 hits may be unused or only triggered outside the log window"
        ws5["A2"].font = Font(name="Calibri", size=10, italic=True); ws5.merge_cells("A2:D2")
        hdr(ws5, 3, ["#", "Rule Name", "Hit Count", "% of Total"])

        all_ssl_names = sorted(set(r["name"] for r in ssl_rules))
        total_ssl = sum(ssl_hits.values()) or 1
        r = 4
        with_hits = [(name, ssl_hits.get(name, 0)) for name in all_ssl_names]
        with_hits.sort(key=lambda x: -x[1])

        for i, (name, count) in enumerate(with_hits, 1):
            cel(ws5, r, 1, i); cel(ws5, r, 2, name); cel(ws5, r, 3, count)
            cel(ws5, r, 4, f"{count/total_ssl*100:.2f}%")
            if count == 0: ws5.cell(row=r, column=3).fill = LIGHT_RED
            r += 1

        log_only = set(ssl_hits.keys()) - set(all_ssl_names)
        if log_only:
            r += 1
            ws5.cell(row=r, column=1, value="Rules in logs but not in current policy").font = BOLD
            ws5.merge_cells(f"A{r}:D{r}"); r += 1
            for name in sorted(log_only):
                cel(ws5, r, 2, name, fill=LIGHT_YELLOW); cel(ws5, r, 3, ssl_hits[name]); r += 1

        set_widths(ws5, [5, 42, 12, 10])
        ws5.auto_filter.ref = f"A3:D{r-1}"; ws5.freeze_panes = "A4"

    # --- TAB: Timeline ---
    ws_tl = wb.create_sheet("Modification Timeline")
    ws_tl["A1"] = "When rules were last modified"
    ws_tl["A1"].font = SUBTITLE
    hdr(ws_tl, 3, ["Month", "URL Rules Modified", "SSL Rules Modified"])
    all_months = sorted(set(list(url_findings["timeline"]["months"].keys()) + list(ssl_findings["timeline"].keys())))
    r = 4
    for month in all_months:
        cel(ws_tl, r, 1, month)
        cel(ws_tl, r, 2, url_findings["timeline"]["months"].get(month, 0))
        cel(ws_tl, r, 3, ssl_findings["timeline"].get(month, 0)); r += 1
    set_widths(ws_tl, [15, 20, 20])

    # --- TAB: Location Distribution ---
    ws_loc = wb.create_sheet("Location Distribution")
    hdr(ws_loc, 1, ["#", "Location", "URL Rules", "SSL Rules (overlapping)"])
    r = 2
    for i, (loc, count) in enumerate(url_findings["locations"]["counts"], 1):
        cel(ws_loc, r, 1, i); cel(ws_loc, r, 2, loc); cel(ws_loc, r, 3, count)
        ssl_overlap = len(ssl_findings["location_overlap"].get(loc, []))
        cel(ws_loc, r, 4, ssl_overlap if ssl_overlap else ""); r += 1
    set_widths(ws_loc, [5, 55, 10, 18]); ws_loc.freeze_panes = "A2"

    # --- TAB: DECRYPT Coverage ---
    ws_dc = wb.create_sheet("DECRYPT Coverage")
    ws_dc["A1"] = "SSL DECRYPT rules and label breakdown"
    ws_dc["A1"].font = SUBTITLE; ws_dc.merge_cells("A1:E1")
    hdr(ws_dc, 2, ["Rule", "Order", "Labels", "Groups", "Predefined"])
    r = 3
    for dc in ssl_findings["decrypt_coverage"]:
        cel(ws_dc, r, 1, dc["name"]); cel(ws_dc, r, 2, dc["order"])
        cel(ws_dc, r, 3, dc["labels"]); cel(ws_dc, r, 4, dc["groups"])
        cel(ws_dc, r, 5, "Yes" if dc["predefined"] else ""); r += 1
    r += 1
    ws_dc.cell(row=r, column=1, value="BYPASS vs DECRYPT by Label").font = SUBTITLE; r += 1
    hdr(ws_dc, r, ["Label", "DECRYPT", "BYPASS", "Skip All Policies", "Evaluate Other"]); r += 1
    for label in sorted(ssl_findings["label_stats"].keys()):
        s = ssl_findings["label_stats"][label]
        cel(ws_dc, r, 1, label); cel(ws_dc, r, 2, s["decrypt"])
        cel(ws_dc, r, 3, s["bypass"]); cel(ws_dc, r, 4, s["bypass_all"])
        cel(ws_dc, r, 5, s["bypass_eval"]); r += 1
    set_widths(ws_dc, [35, 8, 20, 18, 16])

    return wb


def build_docx(url_findings, ssl_findings, url_hits, ssl_hits, url_rules, ssl_rules):
    doc = DocxDocument()
    style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    doc.add_heading("ZIA URL Filtering & SSL Inspection", level=0)
    doc.add_heading("Deep Analysis Report", level=1)
    doc.add_paragraph("Network Security Engineering — Security Control Management")

    doc.add_heading("URL Filtering — Detailed Findings", level=1)

    doc.add_heading("Rule Shadowing", level=2)
    n = len(url_findings["shadows"])
    doc.add_paragraph(
        f'{n} rule shadowing instance{"s" if n != 1 else ""} identified where an ALLOW rule is '
        f'evaluated before a BLOCK rule with overlapping URL categories. The ALLOW rule may be '
        f'permitting traffic the BLOCK rule was intended to deny.'
    )
    for s in url_findings["shadows"]:
        doc.add_paragraph(
            f'ALLOW "{s["allow_name"]}" (order {s["allow_order"]}) shadows '
            f'BLOCK "{s["block_name"]}" (order {s["block_order"]}) — {s["overlap_count"]} overlapping categories.',
            style="List Bullet"
        )

    doc.add_heading("Merge Opportunities", level=2)
    doc.add_paragraph(
        f'{len(url_findings["merges"])} groups of rules share identical scoping and could be consolidated.'
    )

    if url_hits:
        doc.add_heading("URL Filtering Policy Hit Count", level=2)
        total = len(url_rules)
        with_hits = sum(1 for r in url_rules if r["name"] in url_hits)
        zero = total - with_hits
        doc.add_paragraph(
            f'Web Insights logs were analyzed for URL filtering policy hit counts. '
            f'Of {total} rules, {with_hits} had at least one hit in the log period. '
            f'{zero} rules had 0 hits and may be candidates for review or removal. '
            f'Note that 0-hit rules may still be needed for edge cases not captured in the log window.'
        )

    doc.add_heading("SSL Inspection — Detailed Findings", level=1)

    doc.add_heading("DECRYPT Coverage by Label", level=2)
    for label, stats in sorted(ssl_findings["label_stats"].items()):
        ratio = stats["decrypt"] / max(stats["decrypt"] + stats["bypass"], 1) * 100
        doc.add_paragraph(
            f'{label}: {stats["decrypt"]} DECRYPT, {stats["bypass"]} BYPASS '
            f'({ratio:.0f}% inspection). {stats["bypass_all"]} skip all policies, '
            f'{stats["bypass_eval"]} evaluate other policies.',
            style="List Bullet"
        )

    if ssl_hits:
        doc.add_heading("SSL Policy Hit Count", level=2)
        total = len(ssl_rules)
        with_hits = sum(1 for r in ssl_rules if r["name"] in ssl_hits)
        zero = total - with_hits
        doc.add_paragraph(
            f'Of {total} SSL rules, {with_hits} had hits. {zero} rules had 0 hits in the log period.'
        )

    doc.add_heading("Location Overlap", level=2)
    doc.add_paragraph(
        f'{len(ssl_findings["location_overlap"])} server locations have multiple SSL bypass rules, '
        f'representing consolidation opportunities.'
    )

    return doc


def main():
    if len(sys.argv) < 3:
        print(__doc__); sys.exit(1)

    url_path, ssl_path = sys.argv[1], sys.argv[2]
    folder = os.path.dirname(os.path.abspath(url_path)) or os.getcwd()

    for path in [url_path, ssl_path]:
        try: open(path, "r").close(); print(f"  OK: {path}")
        except Exception as e: print(f"ERROR: {path}: {e}"); sys.exit(1)

    print("\nStep 1: Loading policies...")
    with open(url_path, "r", encoding="utf-8") as f: url_rules = json.load(f)
    with open(ssl_path, "r", encoding="utf-8") as f: ssl_rules = json.load(f)
    print(f"  URL filtering: {len(url_rules)} rules")
    print(f"  SSL inspection: {len(ssl_rules)} rules")

    cat_map = load_cat_map(folder)

    print("\nStep 2: Loading Web Insights logs (if present)...")
    url_hits, ssl_hits = load_web_logs(folder)
    if not url_hits and not ssl_hits:
        print("  No hit count data found. Place *_WEB_log.csv files in the same folder to include.")

    print("\nStep 3: Analyzing URL filtering...")
    url_findings = analyze_url_rules(url_rules, cat_map)
    print(f"  Shadowing: {len(url_findings['shadows'])}, Merges: {len(url_findings['merges'])}")

    print("\nStep 4: Analyzing SSL inspection...")
    ssl_findings = analyze_ssl_rules(ssl_rules)
    print(f"  DECRYPT rules: {len(ssl_findings['decrypt_coverage'])}, Location overlaps: {len(ssl_findings['location_overlap'])}")

    print("\nStep 5: Building Excel...")
    wb = build_excel(url_findings, ssl_findings, url_rules, ssl_rules, url_hits, ssl_hits, cat_map)
    xlsx = os.path.join(os.getcwd(), "ZIA_Policy_Deep_Analysis.xlsx")
    wb.save(xlsx); print(f"  Saved: {xlsx}")

    if HAS_DOCX:
        print("\nStep 6: Building Word doc...")
        doc = build_docx(url_findings, ssl_findings, url_hits, ssl_hits, url_rules, ssl_rules)
        docx = os.path.join(os.getcwd(), "ZIA_Policy_Deep_Analysis.docx")
        doc.save(docx); print(f"  Saved: {docx}")

    print("\nDone!")

if __name__ == "__main__":
    main()
